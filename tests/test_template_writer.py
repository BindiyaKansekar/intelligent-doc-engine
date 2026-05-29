"""
tests/test_template_writer.py
=============================
Unit tests for src/template_writer.py.
"""

import shutil
from pathlib import Path

import pytest
from docx import Document

from src.template_writer import write_document


TEMPLATE_PATH = "templates/AUD_template.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _all_text(doc: Document) -> str:
    """Collect all paragraph text from the document into one string."""
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text)
    return "\n".join(parts)


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------

class TestWriteDocument:
    def test_placeholders_are_replaced(self, tmp_path: Path) -> None:
        out = tmp_path / "AUD_v0.1.0.docx"
        replacements = {
            "PROJECT_NAME": "TestProj",
            "VERSION": "0.1.0",
            "GENERATED_DATE": "2026-01-01",
            "OVERVIEW": "This is the overview.",
            "FILE_INVENTORY": "file.py | Python | main module",
            "DATA_FLOWS": "Input -> Process -> Output",
            "DEPENDENCIES": "anthropic, python-docx",
            "CONFIGURATION": "ANTHROPIC_API_KEY: string",
            "KNOWN_ISSUES": "None known.",
        }
        result_path = write_document(TEMPLATE_PATH, str(out), replacements)
        assert result_path.exists()

        doc = Document(str(result_path))
        full_text = _all_text(doc)

        assert "TestProj" in full_text
        assert "0.1.0" in full_text
        assert "This is the overview." in full_text
        # No unreplaced placeholders should remain for keys we provided
        assert "{{OVERVIEW}}" not in full_text
        assert "{{PROJECT_NAME}}" not in full_text

    def test_output_directory_is_created(self, tmp_path: Path) -> None:
        nested = tmp_path / "deep" / "nested" / "dir"
        out = nested / "AUD_test.docx"
        write_document(TEMPLATE_PATH, str(out), {"PROJECT_NAME": "X", "VERSION": "0.0.1"})
        assert out.exists()

    def test_missing_template_raises(self, tmp_path: Path) -> None:
        with pytest.raises(FileNotFoundError):
            write_document(
                str(tmp_path / "no_template.docx"),
                str(tmp_path / "out.docx"),
                {},
            )

    def test_master_template_not_modified(self, tmp_path: Path) -> None:
        """Writing a filled doc must not alter the source template."""
        import hashlib

        def sha(path: str) -> str:
            return hashlib.sha256(Path(path).read_bytes()).hexdigest()

        before = sha(TEMPLATE_PATH)
        write_document(TEMPLATE_PATH, str(tmp_path / "out.docx"), {"PROJECT_NAME": "Y"})
        assert sha(TEMPLATE_PATH) == before
