"""
_make_template.py
-----------------
One-off helper script: generates templates/AUD_template.docx with all
{{PLACEHOLDER}} fields.  Run once during project setup, then delete or
keep as a maintenance utility.

Usage:
    python _make_template.py
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from pathlib import Path


def add_section_heading(doc: Document, title: str) -> None:
    h = doc.add_heading(title, level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_placeholder_paragraph(doc: Document, placeholder: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"{{{{{placeholder}}}}}")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)  # light grey hint colour
    run.font.italic = True


def build_template(output_path: str) -> None:
    doc = Document()

    # ------------------------------------------------------------------ Title
    title = doc.add_heading("Application Understanding Document (AUD)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Project: ").bold = True
    meta.add_run("{{PROJECT_NAME}}\n")
    meta.add_run("Version: ").bold = True
    meta.add_run("{{VERSION}}   ")
    meta.add_run("Generated: ").bold = True
    meta.add_run("{{GENERATED_DATE}}")

    doc.add_paragraph()  # spacer

    # --------------------------------------------------- 1. Project Overview
    add_section_heading(doc, "1. Project Overview")
    doc.add_paragraph(
        "High-level summary of the project's purpose, scope, and key architectural decisions."
    )
    add_placeholder_paragraph(doc, "OVERVIEW")
    doc.add_paragraph()

    # --------------------------------------------------- 2. File Inventory
    add_section_heading(doc, "2. File Inventory")
    doc.add_paragraph(
        "Catalogue of all tracked source files, their types, and their purpose."
    )
    add_placeholder_paragraph(doc, "FILE_INVENTORY")
    doc.add_paragraph()

    # --------------------------------------------------- 3. Data Flows
    add_section_heading(doc, "3. Data Flows")
    doc.add_paragraph(
        "Description of how data moves through the system: inputs, processing steps, and outputs."
    )
    add_placeholder_paragraph(doc, "DATA_FLOWS")
    doc.add_paragraph()

    # --------------------------------------------------- 4. Dependencies
    add_section_heading(doc, "4. Dependencies")
    doc.add_paragraph(
        "External libraries, services, APIs, and databases that the project depends on."
    )
    add_placeholder_paragraph(doc, "DEPENDENCIES")
    doc.add_paragraph()

    # --------------------------------------------------- 5. Configuration
    add_section_heading(doc, "5. Configuration")
    doc.add_paragraph(
        "All configuration keys, environment variables, and their expected values or data types."
    )
    add_placeholder_paragraph(doc, "CONFIGURATION")
    doc.add_paragraph()

    # --------------------------------------------------- 6. Known Issues
    add_section_heading(doc, "6. Known Issues")
    doc.add_paragraph(
        "TODOs, deprecated patterns, potential bugs, and technical debt identified in the codebase."
    )
    add_placeholder_paragraph(doc, "KNOWN_ISSUES")
    doc.add_paragraph()

    # ----------------------------------------------------------------- Footer
    section = doc.sections[0]
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("AUTO-GENERATED — DO NOT EDIT MANUALLY  |  ").font.size = Pt(8)
    run = footer_para.add_run("{{PROJECT_NAME}} v{{VERSION}}")
    run.font.size = Pt(8)

    # ------------------------------------------------------------------ Save
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    print(f"Template written to: {out.resolve()}")


if __name__ == "__main__":
    build_template("templates/AUD_template.docx")
