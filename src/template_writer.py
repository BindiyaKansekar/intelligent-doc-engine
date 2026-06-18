"""
template_writer.py
==================
Word document (.docx) placeholder injection for the Intelligent Documentation Engine.

Responsibilities:
  - Load the AUD_template.docx master template.
  - Replace every {{PLACEHOLDER}} token in paragraphs, table cells, and headers/
    footers with the corresponding generated text.
  - Insert formatted tables for pipe-delimited section content.
  - Insert images for rendered diagrams.
  - Save the filled document to the appropriate versioned output path.
"""

from __future__ import annotations

import logging
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

logger = logging.getLogger(__name__)

# Pattern that matches {{SOME_KEY}} tokens (uppercase letters and underscores only)
_PLACEHOLDER_RE = re.compile(r"\{\{([A-Z_]+)\}\}")


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def _parse_pipe_table(text: str) -> tuple[list[str], list[list[str]]] | None:
    """Parse pipe-delimited table text into (headers, rows).

    Returns None if the text doesn't look like a table.
    """
    lines = [ln.strip() for ln in text.strip().splitlines() if ln.strip() and "|" in ln]
    if len(lines) < 2:
        return None

    def split_row(line: str) -> list[str]:
        cells = [c.strip() for c in line.split("|")]
        if cells and not cells[0]:
            cells = cells[1:]
        if cells and not cells[-1]:
            cells = cells[:-1]
        return cells

    headers = split_row(lines[0])
    if not headers:
        return None

    rows: list[list[str]] = []
    for line in lines[1:]:
        if set(line.replace("|", "").replace("-", "").strip()) <= {" ", ""}:
            continue
        cells = split_row(line)
        if cells:
            rows.append(cells)
    return (headers, rows) if rows else None


def _insert_table(doc: Document, paragraph, headers: list[str], rows: list[list[str]]) -> None:
    """Replace *paragraph* with a formatted Word table."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)

    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row[:num_cols]):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    paragraph._element.addnext(table._tbl)
    paragraph.text = ""


def _insert_image(doc: Document, paragraph, image_path: str, width_inches: float = 5.5) -> None:
    """Replace *paragraph* text with an embedded image."""
    paragraph.text = ""
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph: object, replacements: dict[str, str]) -> None:
    """Replace all placeholder tokens in *paragraph* in-place."""
    full_text = "".join(run.text for run in paragraph.runs)

    if not _PLACEHOLDER_RE.search(full_text):
        return

    def substitute(match: re.Match) -> str:
        key = match.group(1)
        value = replacements.get(key, match.group(0))
        if not replacements.get(key):
            logger.debug("No replacement value for placeholder {{%s}} — leaving as-is", key)
        return value

    new_text = _PLACEHOLDER_RE.sub(substitute, full_text)

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_document(doc: object, replacements: dict[str, str]) -> None:
    """Walk every text-bearing element in *doc* and apply *replacements*."""
    for para in doc.paragraphs:
        _replace_in_paragraph(para, replacements)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, replacements)

    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf is not None:
                for para in hf.paragraphs:
                    _replace_in_paragraph(para, replacements)


def _process_tables_and_images(
    doc: Document,
    replacements: dict[str, str],
    table_sections: list[str],
    image_sections: dict[str, str],
) -> set[str]:
    """Process table and image sections before plain text replacement.

    Returns set of keys that were handled (should be excluded from text replacement).
    """
    handled: set[str] = set()

    for para in doc.paragraphs:
        full_text = "".join(run.text for run in para.runs)
        match = _PLACEHOLDER_RE.search(full_text)
        if not match:
            continue
        key = match.group(1)

        if key in image_sections:
            img_path = image_sections[key]
            if Path(img_path).exists():
                _insert_image(doc, para, img_path)
                handled.add(key)
                logger.info("Inserted image for {{%s}} from %s", key, img_path)
            continue

        if key in table_sections and key in replacements:
            parsed = _parse_pipe_table(replacements[key])
            if parsed:
                headers, rows = parsed
                _insert_table(doc, para, headers, rows)
                handled.add(key)
                logger.info("Inserted table for {{%s}} (%d rows)", key, len(rows))


    return handled


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

_EXTENDED_SECTION_TITLES: dict[str, str] = {
    "IMPLEMENTATION_DETAILS": "Implementation Details",
    "DATA_MODEL": "Data Model",
    "API_REFERENCE": "API Reference",
}

_EXTENDED_TABLE_SECTIONS = {"DATA_MODEL", "API_REFERENCE"}


def append_extended_sections(
    doc_path: str,
    sections: dict[str, str],
    image_sections: dict[str, str] | None = None,
) -> None:
    """Append Implementation Details, Data Model, and API Reference to the document.

    Also embeds the architecture diagram when *image_sections* contains
    ``ARCHITECTURE_DIAGRAM``, placing it as the first item in the block.

    Args:
        doc_path:      Path to the .docx file to extend.
        sections:      Mapping of extended section key → generated text.
        image_sections: Mapping of section key → rendered image path.
                        Pass the builder's ``image_sections`` dict here so that
                        ``ARCHITECTURE_DIAGRAM`` is embedded directly in the document.
    """
    images = image_sections or {}
    if not sections and "ARCHITECTURE_DIAGRAM" not in images:
        return

    out = Path(doc_path)
    if not out.exists():
        logger.warning(
            "[TemplateWriter] Cannot append extended sections — file not found: %s", out
        )
        return

    doc: Document = Document(str(out))
    doc.add_page_break()
    heading = doc.add_heading("Extended Documentation", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Embed architecture diagram first if available
    arch_img = images.get("ARCHITECTURE_DIAGRAM")
    if arch_img and Path(arch_img).exists():
        doc.add_heading("Architecture Diagram", level=2)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        run.add_picture(arch_img, width=Inches(5.5))
        logger.info("[TemplateWriter] Architecture diagram embedded from %s", arch_img)

    canonical_order = [k for k in _EXTENDED_SECTION_TITLES if k in sections]

    for key in canonical_order:
        content = sections[key]
        title = _EXTENDED_SECTION_TITLES[key]
        doc.add_heading(title, level=2)

        if key in _EXTENDED_TABLE_SECTIONS:
            # DATA_MODEL may contain multiple tables separated by a blank line
            # between model blocks — split and render each independently
            blocks = _split_table_blocks(content)
            for block in blocks:
                parsed = _parse_pipe_table(block)
                if parsed:
                    headers, rows = parsed
                    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                    table.style = "Table Grid"
                    for i, h in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = h
                        for run in cell.paragraphs[0].runs:
                            run.bold = True
                            run.font.size = Pt(9)
                    for r_idx, row in enumerate(rows):
                        for c_idx, val in enumerate(row[:len(headers)]):
                            cell = table.rows[r_idx + 1].cells[c_idx]
                            cell.text = val
                            for run in cell.paragraphs[0].runs:
                                run.font.size = Pt(9)
                    doc.add_paragraph()  # spacing after table
                    logger.info(
                        "[TemplateWriter] Table for '%s' inserted (%d rows)", key, len(rows)
                    )
                else:
                    doc.add_paragraph(block.strip())
        else:
            for line in content.splitlines():
                stripped = line.strip()
                para = doc.add_paragraph()
                run = para.add_run(stripped)
                run.font.size = Pt(10)

    doc.save(str(out))
    logger.info("[TemplateWriter] Extended sections appended to %s", out)


def _split_table_blocks(text: str) -> list[str]:
    """Split text into blocks that may each contain a separate pipe table.

    Handles the DATA_MODEL pattern where Claude emits multiple tables
    (one per model) separated by blank lines.
    """
    blocks: list[str] = []
    current: list[str] = []
    for line in text.splitlines():
        if not line.strip() and current:
            block = "\n".join(current)
            if "|" in block:
                blocks.append(block)
            current = []
        else:
            current.append(line)
    if current:
        block = "\n".join(current)
        if "|" in block:
            blocks.append(block)
    return blocks if blocks else [text]


_INTEGRATION_SECTION_TITLES: dict[str, str] = {
    "INTEGRATION_OVERVIEW": "Integration Overview",
    "SEQUENCE_DIAGRAM": "Sequence / Flow Diagram",
    "API_FIELD_MAPPING": "API & Field Mapping",
    "ERROR_RETRY_BEHAVIOUR": "Error Handling & Retry Behaviour",
}

_INTEGRATION_TABLE_SECTIONS = {"API_FIELD_MAPPING"}
_INTEGRATION_IMAGE_SECTIONS = {"SEQUENCE_DIAGRAM", "ARCHITECTURE_DIAGRAM"}


def append_integration_sections(
    doc_path: str,
    sections: dict[str, str],
    image_sections: dict[str, str] | None = None,
) -> None:
    """Append a 'New Integration Details' block to an existing document.

    Opens the already-written .docx at *doc_path*, adds a page break, then
    writes each integration section with a sub-heading and formatted content
    (embedded image, table, or plain text).  Saves in-place.

    Args:
        doc_path:      Path to the .docx file to extend.
        sections:      Mapping of integration section key → generated text.
        image_sections: Mapping of section key → rendered image path.
    """
    if not sections:
        return

    images = image_sections or {}
    out = Path(doc_path)
    if not out.exists():
        logger.warning("[TemplateWriter] Cannot append integration sections — file not found: %s", out)
        return

    doc: Document = Document(str(out))

    # Page break before integration details
    doc.add_page_break()

    heading = doc.add_heading("New Integration Details", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

    canonical_order = [k for k in _INTEGRATION_SECTION_TITLES if k in sections]

    for key in canonical_order:
        content = sections[key]
        title = _INTEGRATION_SECTION_TITLES[key]

        doc.add_heading(title, level=2)

        if key in images and Path(images[key]).exists():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(images[key], width=Inches(5.5))
            logger.info("[TemplateWriter] Embedded image for '%s' from %s", key, images[key])
            continue

        if key in _INTEGRATION_TABLE_SECTIONS:
            parsed = _parse_pipe_table(content)
            if parsed:
                headers, rows = parsed
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = "Table Grid"
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(9)
                for r_idx, row in enumerate(rows):
                    for c_idx, val in enumerate(row[:len(headers)]):
                        cell = table.rows[r_idx + 1].cells[c_idx]
                        cell.text = val
                        for run in cell.paragraphs[0].runs:
                            run.font.size = Pt(9)
                logger.info("[TemplateWriter] Inserted table for '%s' (%d rows)", key, len(rows))
                continue

        # Plain text / mermaid fallback — strip raw mermaid fences for readability
        text = re.sub(r"```mermaid\s*", "", content)
        text = re.sub(r"```\s*", "", text).strip()
        for line in text.splitlines():
            para = doc.add_paragraph(line if line.strip() else "")
            for run in para.runs:
                run.font.size = Pt(10)

    doc.save(str(out))
    logger.info("[TemplateWriter] Integration sections appended to %s", out)


def append_additional_points(doc_path: str, points: list[str]) -> None:
    """Append a 'Notes & Observations' section with bullet points to an existing document.

    Args:
        doc_path: Path to the .docx file to extend.
        points:   List of observation strings to render as bullet points.
    """
    if not points:
        return

    out = Path(doc_path)
    if not out.exists():
        logger.warning(
            "[TemplateWriter] Cannot append additional points — file not found: %s", out
        )
        return

    doc: Document = Document(str(out))
    doc.add_heading("Notes & Observations", level=2)

    for point in points:
        para = doc.add_paragraph()
        run = para.add_run(f"•  {point.strip()}")
        run.font.size = Pt(10)

    doc.save(str(out))
    logger.info(
        "[TemplateWriter] %d observation(s) appended to %s", len(points), out
    )


def write_document(
    template_path: str,
    output_path: str,
    replacements: dict[str, str],
    table_sections: list[str] | None = None,
    image_sections: dict[str, str] | None = None,
) -> Path:
    """Fill the Word template and write the result to *output_path*.

    Args:
        template_path:  Path to the ``AUD_template.docx`` master file.
        output_path:    Destination path for the filled document.
        replacements:   Mapping of placeholder key → replacement text.
        table_sections: Section keys to render as formatted tables.
        image_sections: Mapping of section key → image file path.

    Returns:
        The resolved :class:`pathlib.Path` of the written document.
    """
    tmpl = Path(template_path)
    if not tmpl.exists():
        raise FileNotFoundError(f"Template not found: {tmpl}")

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    shutil.copy2(tmpl, out)
    logger.debug("Copied template %s → %s", tmpl, out)

    doc: Document = Document(str(out))

    handled: set[str] = set()
    if table_sections or image_sections:
        handled = _process_tables_and_images(
            doc,
            replacements,
            table_sections or [],
            image_sections or {},
        )

    text_replacements = {k: v for k, v in replacements.items() if k not in handled}
    _replace_in_document(doc, text_replacements)
    doc.save(str(out))

    logger.info("Document written to %s", out)
    return out.resolve()
