"""
analytics_docx_writer.py
========================
Two entry points:

  write_sql_docx() — 12-section structured Word document built directly from
                     parsed SQLFileInfo + LineageGraph (navy/teal/slate scheme,
                     layer-colour-coded bands, column lineage grouped by target).

  write_docx()     — legacy Markdown → Word converter used for ADF / Azure
                     Functions / generic output.
"""
from __future__ import annotations

import re
import tempfile
import logging
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from .mermaid_renderer import render_mermaid_to_png

logger = logging.getLogger(__name__)

# ── Brand colours ─────────────────────────────────────────────────────────────
_NAVY  = RGBColor(0x1F, 0x49, 0x7D)
_TEAL  = RGBColor(0x21, 0x73, 0x46)
_SLATE = RGBColor(0x44, 0x54, 0x6A)
_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
_LIGHT = RGBColor(0xF2, 0xF2, 0xF2)

_NAVY_HEX  = "1F497D"
_TEAL_HEX  = "217346"
_SLATE_HEX = "44546A"

# Layer accent colours (dark band header / light row fill)
_LAYER_DARK_HEX = {
    "raw":     "C00000",   # dark red
    "stage":   "44546A",   # slate
    "silver":  "44546A",   # slate
    "mart":    "1D6045",   # dark teal
    "gold":    "1D6045",
    "unknown": "595959",
}
_LAYER_LIGHT_HEX = {
    "raw":     "FCE4D6",   # light rose
    "stage":   "DAE3F3",   # light slate-blue
    "silver":  "DAE3F3",
    "mart":    "E2EFDA",   # light teal-green
    "gold":    "E2EFDA",
    "unknown": "F2F2F2",
}
_LAYER_ORDER = ["raw", "stage", "silver", "mart", "gold", "unknown"]
_LAYER_LABEL = {
    "raw": "RAW", "stage": "STAGE", "silver": "SILVER",
    "mart": "MART", "gold": "GOLD/MART", "unknown": "Source",
}

# Sections grouped as RAW / SILVER / GOLD-MART
_SECTION_GROUPS = [
    (["raw"],          "Layer Summary (RAW)",      "raw"),
    (["stage", "silver"], "Layer Summary (SILVER)", "silver"),
    (["mart", "gold"], "Layer Summary (GOLD/MART)", "gold"),
]


# ══════════════════════════════════════════════════════════════════════════════
# Public API
# ══════════════════════════════════════════════════════════════════════════════

def write_sql_docx(
    sql_infos: list,
    graph,
    output_path: str,
    title: str = "Repository Documentation",
    pr_description: str = "",
    scope: str = "PR changes only",
) -> Path:
    """Build a 12-section structured .docx from parsed SQL + lineage data."""
    doc = Document()
    _setup_styles(doc)

    # ── 1. Cover page ──────────────────────────────────────────────────────
    _add_cover(doc, title, scope, pr_description)

    # ── 2. Table of Contents ───────────────────────────────────────────────
    doc.add_page_break()
    _h1(doc, "Table of Contents")
    _add_toc_field(doc)

    # ── 3. Repository Overview ─────────────────────────────────────────────
    doc.add_page_break()
    _h1(doc, "Repository Overview")
    _add_overview(doc, sql_infos, pr_description, scope)

    # ── 4. Data Architecture ───────────────────────────────────────────────
    _h1(doc, "Data Architecture")
    _add_architecture(doc, sql_infos, graph)

    # ── 5–7. Layer Summaries ───────────────────────────────────────────────
    by_layer: dict[str, list] = defaultdict(list)
    for f in sql_infos:
        by_layer[f.layer].append(f)

    for layer_keys, section_name, color_key in _SECTION_GROUPS:
        files = [f for k in layer_keys for f in by_layer.get(k, [])]
        if files:
            _add_layer_summary(doc, section_name, files, color_key)

    # ── 8. Object Inventory ────────────────────────────────────────────────
    if sql_infos:
        _h1(doc, "Object Inventory")
        _add_object_inventory(doc, sql_infos)

    # ── 9. Data Lineage ────────────────────────────────────────────────────
    if graph and getattr(graph, "edges", None):
        _h1(doc, "Data Lineage")
        _add_lineage_diagram(doc, graph, title)

    # ── 10. Column-Level Lineage ───────────────────────────────────────────
    if graph and getattr(graph, "edge_columns", None):
        _h1(doc, "Column-Level Lineage")
        _add_column_lineage(doc, graph)

    # ── 11. Column Inventory ───────────────────────────────────────────────
    files_with_cols = [f for f in sql_infos if getattr(f, "columns", None)]
    if files_with_cols:
        _h1(doc, "Column Inventory")
        _add_column_inventory(doc, files_with_cols)

    # ── 12. Change Summary ─────────────────────────────────────────────────
    _h1(doc, "Change Summary")
    _add_change_summary(doc, sql_infos, scope)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("Structured Word document written → %s", out)
    return out.resolve()


# ══════════════════════════════════════════════════════════════════════════════
# Section builders
# ══════════════════════════════════════════════════════════════════════════════

def _add_cover(doc: Document, title: str, scope: str, description: str) -> None:
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = _NAVY

    doc.add_paragraph()

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Intelligent Doc Engine  ·  Analytics Documentation")
    r.font.size = Pt(14)
    r.font.color.rgb = _TEAL

    doc.add_paragraph()

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts = datetime.now(timezone.utc).strftime("%d %B %Y  %H:%M UTC")
    r2 = meta.add_run(f"{ts}   |   Scope: {scope}")
    r2.font.size = Pt(10)
    r2.font.color.rgb = _SLATE

    if description:
        doc.add_paragraph()
        desc_p = doc.add_paragraph()
        desc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = desc_p.add_run(description[:200])
        dr.font.size = Pt(10)
        dr.font.italic = True
        dr.font.color.rgb = _SLATE

    doc.add_paragraph()
    rule = doc.add_paragraph("─" * 80)
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rule.runs[0].font.color.rgb = _NAVY


def _add_toc_field(doc: Document) -> None:
    note = doc.add_paragraph()
    nr = note.add_run(
        "Right-click this area in Word and select 'Update Field' to populate the table of contents."
    )
    nr.font.size = Pt(9)
    nr.font.italic = True
    nr.font.color.rgb = _SLATE

    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr_run = para.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run._r.append(instr)

    sep_run = para.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    ph_run = para.add_run("Table of Contents will appear here after updating the field.")
    ph_run.font.color.rgb = _SLATE
    ph_run.font.italic = True

    end_run = para.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def _add_overview(doc: Document, sql_infos: list, description: str, scope: str) -> None:
    layers_present = sorted(
        {f.layer for f in sql_infos if f.layer != "unknown"},
        key=lambda l: _LAYER_ORDER.index(l) if l in _LAYER_ORDER else 99,
    )
    obj_types = sorted({f.object_type for f in sql_infos if f.object_type != "UNKNOWN"})
    total_cols = sum(len(getattr(f, "columns", [])) for f in sql_infos)

    rows = [
        ("Total SQL files", str(len(sql_infos))),
        ("Layers present", "  →  ".join(_LAYER_LABEL.get(l, l.upper()) for l in layers_present) or "N/A"),
        ("Object types", ", ".join(obj_types) or "N/A"),
        ("Total columns tracked", str(total_cols)),
        ("Documentation scope", scope),
    ]
    _kv_table(doc, rows)

    if description:
        doc.add_paragraph()
        p = doc.add_paragraph(description)
        p.runs[0].font.size = Pt(10)


def _add_architecture(doc: Document, sql_infos: list, graph) -> None:
    by_layer: dict[str, list] = defaultdict(list)
    for f in sql_infos:
        by_layer[f.layer].append(f)

    layers_present = [l for l in _LAYER_ORDER if by_layer.get(l)]
    flow = "  →  ".join(_LAYER_LABEL.get(l, l.upper()) for l in layers_present)
    p = doc.add_paragraph()
    r = p.add_run(f"Data flow:  {flow}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = _NAVY

    doc.add_paragraph()

    desc_rows = []
    for layer in layers_present:
        files = by_layer[layer]
        label = _LAYER_LABEL.get(layer, layer.upper())
        src_set: set[str] = set()
        for f in files:
            src_set.update(getattr(f, "sources", []))
        desc = (
            f"{len(files)} object(s). "
            + (f"Reads from: {', '.join(sorted(src_set)[:5])}." if src_set else "Primary ingestion layer.")
        )
        desc_rows.append((label, desc))

    _kv_table(doc, desc_rows, key_width=Inches(1.2))
    doc.add_paragraph()


def _add_layer_summary(doc: Document, section_name: str, files: list, color_key: str) -> None:
    _h1(doc, section_name)

    dark_hex = _LAYER_DARK_HEX.get(color_key, "595959")
    light_hex = _LAYER_LIGHT_HEX.get(color_key, "F2F2F2")

    headers = ["File", "Object Name", "Type", "Columns", "Load Strategy", "Source Tables"]
    col_rows = []
    for f in files:
        fname = f.path.replace("\\", "/").split("/")[-1]
        obj = getattr(f, "primary_target", None) or (f.targets[0] if getattr(f, "targets", None) else None)
        obj_name = str(obj) if obj else getattr(f, "object_name", None) or "—"
        ncols = len(getattr(f, "columns", []))
        strategy = getattr(f, "load_strategy", None) or "—"
        sources = getattr(f, "sources", [])
        src_str = ", ".join(str(s).split(".")[-1] for s in sources[:4])
        if len(sources) > 4:
            src_str += f" +{len(sources) - 4} more"
        col_rows.append([fname, obj_name, f.object_type, str(ncols), strategy, src_str or "—"])

    _styled_table(doc, headers, col_rows, header_hex=dark_hex, alt_hex=light_hex)
    doc.add_paragraph()


def _add_object_inventory(doc: Document, sql_infos: list) -> None:
    headers = ["Layer", "Object Name", "Type", "File", "Columns", "Load Strategy"]
    rows = []
    for layer in _LAYER_ORDER:
        for f in [x for x in sql_infos if x.layer == layer]:
            obj = getattr(f, "primary_target", None) or (f.targets[0] if getattr(f, "targets", None) else None)
            obj_name = str(obj) if obj else getattr(f, "object_name", None) or "—"
            fname = f.path.replace("\\", "/").split("/")[-1]
            rows.append([
                _LAYER_LABEL.get(layer, layer.upper()),
                obj_name,
                f.object_type,
                fname,
                str(len(getattr(f, "columns", []))),
                getattr(f, "load_strategy", None) or "—",
            ])

    _styled_table(doc, headers, rows, header_hex=_NAVY_HEX, alt_hex="EBF0F7")
    doc.add_paragraph()


def _add_lineage_diagram(doc: Document, graph, title: str) -> None:
    diagrams = graph.to_layer_diagrams(title_prefix=f"Lineage — {title}")
    for _, mermaid_code in diagrams:
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            png_path = tmp.name
        result = render_mermaid_to_png(mermaid_code, png_path)
        if result and Path(png_path).exists():
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(png_path, width=Inches(_fit_width(png_path)))
            doc.add_paragraph()
        else:
            p = doc.add_paragraph("Lineage diagram could not be rendered. "
                                  "Install mermaid-cli (npm i -g @mermaid-js/mermaid-cli).")
            p.runs[0].font.italic = True
            p.runs[0].font.color.rgb = _SLATE


def _add_column_lineage(doc: Document, graph) -> None:
    # Group all column mappings by TARGET table
    by_target: dict[str, list[tuple[str, str, str]]] = defaultdict(list)
    for (src, tgt), pairs in sorted(graph.edge_columns.items()):
        for src_col, tgt_col in pairs:
            by_target[tgt].append((src, src_col, tgt_col))

    for tgt in sorted(by_target):
        tgt_label = tgt.split(".")[-1]
        _h2(doc, f"Target: {tgt_label}")

        mappings = sorted(by_target[tgt], key=lambda x: (x[0], x[2]))
        headers = ["Source Table", "Source Column", "Target Column"]
        rows = [
            [src.split(".")[-1], src_col or "(expression)", tgt_col]
            for src, src_col, tgt_col in mappings
        ]
        _styled_table(doc, headers, rows, header_hex=_TEAL_HEX, alt_hex="E8F4EE")
        doc.add_paragraph()


def _add_column_inventory(doc: Document, sql_infos: list) -> None:
    for layer in _LAYER_ORDER:
        layer_files = [f for f in sql_infos if f.layer == layer and getattr(f, "columns", None)]
        if not layer_files:
            continue

        _h2(doc, _LAYER_LABEL.get(layer, layer.upper()))
        dark_hex = _LAYER_DARK_HEX.get(layer, "595959")
        light_hex = _LAYER_LIGHT_HEX.get(layer, "F2F2F2")

        for f in layer_files:
            obj = getattr(f, "primary_target", None) or (f.targets[0] if getattr(f, "targets", None) else None)
            obj_name = str(obj) if obj else getattr(f, "object_name", None) or f.path.split("/")[-1]

            _h3(doc, obj_name)
            pk_set = {c.upper() for c in getattr(f, "pk_columns", [])}
            rows = [
                [c, "PK" if c.upper() in pk_set else ""]
                for c in f.columns
            ]
            _styled_table(doc, ["Column", "Key"], rows, header_hex=dark_hex, alt_hex=light_hex)
            doc.add_paragraph()


def _add_change_summary(doc: Document, sql_infos: list, scope: str) -> None:
    by_layer: dict[str, list] = defaultdict(list)
    for f in sql_infos:
        by_layer[f.layer].append(f)

    p = doc.add_paragraph()
    r = p.add_run(f"Scope: {scope}  ·  {len(sql_infos)} file(s) processed")
    r.font.size = Pt(10)
    r.font.color.rgb = _SLATE
    doc.add_paragraph()

    headers = ["Layer", "File", "Object", "Type", "Columns", "Sources"]
    rows = []
    for layer in _LAYER_ORDER:
        for f in by_layer.get(layer, []):
            obj = getattr(f, "primary_target", None) or (f.targets[0] if getattr(f, "targets", None) else None)
            obj_name = str(obj) if obj else getattr(f, "object_name", None) or "—"
            fname = f.path.replace("\\", "/").split("/")[-1]
            sources = getattr(f, "sources", [])
            rows.append([
                _LAYER_LABEL.get(layer, layer.upper()),
                fname,
                obj_name,
                f.object_type,
                str(len(getattr(f, "columns", []))),
                str(len(sources)),
            ])

    if rows:
        _styled_table(doc, headers, rows, header_hex=_NAVY_HEX, alt_hex="EBF0F7")

    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Generated by Intelligent Doc Engine  ·  {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True
    footer.runs[0].font.color.rgb = _SLATE


# ══════════════════════════════════════════════════════════════════════════════
# Shared primitives
# ══════════════════════════════════════════════════════════════════════════════

def _h1(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.color.rgb = _NAVY


def _h2(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.color.rgb = _TEAL


def _h3(doc: Document, text: str) -> None:
    h = doc.add_heading(text, level=3)
    for run in h.runs:
        run.font.color.rgb = _SLATE


def _kv_table(
    doc: Document,
    rows: list[tuple[str, str]],
    key_width: object = None,
) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for ri, (key, val) in enumerate(rows):
        k_cell = table.rows[ri].cells[0]
        v_cell = table.rows[ri].cells[1]
        k_cell.text = key
        v_cell.text = val
        for run in k_cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _NAVY
        for run in v_cell.paragraphs[0].runs:
            run.font.size = Pt(9)
        if ri % 2 == 0:
            _shade_cell(k_cell, "EBF0F7")
            _shade_cell(v_cell, "EBF0F7")


def _styled_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_hex: str = _NAVY_HEX,
    alt_hex: str = "EBF0F7",
) -> None:
    if not rows and not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = h
        _shade_cell(cell, header_hex)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _WHITE

    # Data rows
    for ri, row_data in enumerate(rows):
        shade = alt_hex if ri % 2 == 1 else None
        for ci, val in enumerate(row_data[: len(headers)]):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val else "—"
            if shade:
                _shade_cell(cell, shade)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _setup_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)


def _fit_width(png_path: str, max_w: float = 6.0, max_h: float = 4.0) -> float:
    import struct
    try:
        with open(png_path, "rb") as f:
            f.read(16)
            w_px = struct.unpack(">I", f.read(4))[0]
            h_px = struct.unpack(">I", f.read(4))[0]
        dpi = 96
        scale = min(1.0, max_w / (w_px / dpi), max_h / (h_px / dpi))
        return max(1.5, (w_px / dpi) * scale)
    except Exception:
        return max_w


# ══════════════════════════════════════════════════════════════════════════════
# Legacy Markdown → Word (ADF / Functions / generic)
# ══════════════════════════════════════════════════════════════════════════════

_MERMAID_BLOCK_RE = re.compile(r"```mermaid\s*\n(.*?)```", re.DOTALL)


def write_docx(markdown: str, output_path: str, title: str = "Analytics Documentation") -> Path:
    """Convert Markdown to a Word document (legacy path for non-SQL repos)."""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _setup_styles(doc)

    block_replacements: dict[str, str] = {}
    for i, m in enumerate(_MERMAID_BLOCK_RE.finditer(markdown)):
        png_path = str(Path(tempfile.gettempdir()) / f"analytics_diagram_{i}.png")
        result = render_mermaid_to_png(m.group(1).strip(), png_path)
        block_replacements[m.group(0)] = f"[[DIAGRAM:{png_path}]]" if result else "[[DIAGRAM_FAILED]]"

    processed = markdown
    for raw_block, replacement in block_replacements.items():
        processed = processed.replace(raw_block, replacement)

    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = _NAVY

    sub = doc.add_paragraph(
        f"Generated by Intelligent Doc Engine  ·  {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = _SLATE
    doc.add_page_break()

    lines = processed.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("# ") and not line.startswith("## "):
            h = doc.add_heading(line[2:].strip(), level=1)
            for run in h.runs:
                run.font.color.rgb = _NAVY
            i += 1; continue

        if line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            for run in h.runs:
                run.font.color.rgb = _TEAL
            i += 1; continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1; continue

        if line.startswith("[[DIAGRAM:") and line.endswith("]]"):
            png_path = line[10:-2]
            if Path(png_path).exists():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(png_path, width=Inches(_fit_width(png_path)))
                doc.add_paragraph()
            else:
                doc.add_paragraph("[Diagram image not available]").runs[0].font.color.rgb = _SLATE
            i += 1; continue

        if line == "[[DIAGRAM_FAILED]]":
            p = doc.add_paragraph("Lineage diagram could not be rendered.")
            p.runs[0].font.color.rgb = _SLATE
            p.runs[0].font.italic = True
            i += 1; continue

        if "|" in line and line.strip().startswith("|"):
            table_lines = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                table_lines.append(lines[i]); i += 1
            _write_md_table(doc, table_lines)
            continue

        if line.startswith("- ") or line.startswith("* "):
            para = doc.add_paragraph(line[2:].strip(), style="List Bullet")
            para.runs[0].font.size = Pt(10)
            i += 1; continue

        if re.match(r"^\d+\.\s", line):
            stripped = re.sub(r"^\d+\.\s", "", line).strip()
            para = doc.add_paragraph(stripped, style="List Number")
            para.runs[0].font.size = Pt(10)
            i += 1; continue

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i]); i += 1
            i += 1
            if code_lines:
                para = doc.add_paragraph("\n".join(code_lines))
                para.runs[0].font.name = "Courier New"
                para.runs[0].font.size = Pt(8)
            continue

        if "**" in line:
            _write_inline_bold(doc, line)
            i += 1; continue

        if line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)
            i += 1; continue

        if line.strip():
            para = doc.add_paragraph(line.strip())
            para.runs[0].font.size = Pt(10)
        i += 1

    doc.save(str(out))
    logger.info("Word document written → %s", out)
    return out.resolve()


def _write_md_table(doc: Document, table_lines: list[str]) -> None:
    def split_row(line: str) -> list[str]:
        return [c.strip() for c in line.split("|") if c.strip()]

    data_lines = [
        ln for ln in table_lines
        if not re.match(r"^\s*\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?\s*$", ln)
    ]
    if not data_lines:
        return

    headers = split_row(data_lines[0])
    rows = [split_row(ln) for ln in data_lines[1:]]
    if not headers:
        return

    _styled_table(doc, headers, rows, header_hex=_NAVY_HEX)
    doc.add_paragraph()


def _write_inline_bold(doc: Document, line: str) -> None:
    para = doc.add_paragraph()
    for idx, part in enumerate(re.split(r"\*\*", line)):
        if not part:
            continue
        run = para.add_run(part)
        run.bold = (idx % 2 == 1)
        run.font.size = Pt(10)
